from __future__ import annotations

import re
from collections import Counter
from statistics import median
from typing import Any, Dict, Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}

ROLE_LABELS = {
    "title": "\u4e2d\u6587\u6807\u9898",
    "authors": "\u4e2d\u6587\u4f5c\u8005",
    "affiliations": "\u4e2d\u6587\u5355\u4f4d",
    "abstract_cn": "\u4e2d\u6587\u6458\u8981",
    "keywords_cn": "\u4e2d\u6587\u5173\u952e\u8bcd",
    "title_en": "\u82f1\u6587\u6807\u9898",
    "authors_en": "\u82f1\u6587\u4f5c\u8005",
    "affiliations_en": "\u82f1\u6587\u5355\u4f4d",
    "abstract_en": "\u82f1\u6587\u6458\u8981",
    "keywords_en": "\u82f1\u6587\u5173\u952e\u8bcd",
    "heading_1": "\u4e00\u7ea7\u6807\u9898",
    "heading_2": "\u4e8c\u7ea7\u6807\u9898",
    "body": "\u6b63\u6587",
    "fund": "\u57fa\u91d1\u9879\u76ee",
    "caption_figure": "\u56fe\u9898",
    "caption_table": "\u8868\u9898",
    "references_title": "\u53c2\u8003\u6587\u732e\u6807\u9898",
    "reference_entry": "\u53c2\u8003\u6587\u732e\u6761\u76ee",
    "author_bio_title": "\u4f5c\u8005\u7b80\u4ecb\u6807\u9898",
    "author_bio_entry": "\u4f5c\u8005\u7b80\u4ecb",
    "header": "\u9875\u7709",
    "footer": "\u9875\u811a",
}

TEXT = {
    "abstract": "\u6458\u8981",
    "keywords": "\u5173\u952e\u8bcd",
    "keywords_alt": "\u5173\u952e\u5b57",
    "fund": "\u57fa\u91d1\u9879\u76ee",
    "references": "\u53c2\u8003\u6587\u732e",
    "received": "\u6536\u7a3f\u65e5\u671f",
    "author_bio": "\u4f5c\u8005\u7b80\u4ecb",
    "intro": "\u5f15\u8a00",
    "conclusion": "\u7ed3\u8bba",
    "figure": "\u56fe",
    "table": "\u8868",
    "figure_caption": "\u56fe\u9898",
    "table_caption": "\u8868\u9898",
}

SKIP_PREFIXES = (
    "doi:",
    TEXT["abstract"],
    TEXT["keywords"],
    TEXT["keywords_alt"],
    "abstract",
    "key words",
    "keywords",
    TEXT["fund"],
    TEXT["references"],
    TEXT["received"],
    TEXT["author_bio"],
)


def _fmt_length(length) -> Optional[float]:
    return round(length.pt, 2) if length is not None else None


def _style_chain(style) -> Iterable:
    current = style
    while current is not None:
        yield current
        current = current.base_style


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\xa0", " ")).strip()


def contains_cjk(text: str) -> bool:
    return any("\u4e00" <= ch <= "\u9fff" for ch in text)


def is_ascii_heavy(text: str) -> bool:
    cleaned = normalize_text(text)
    if not cleaned:
        return False
    ascii_chars = sum(1 for ch in cleaned if ord(ch) < 128 and ch.isprintable())
    return ascii_chars / max(len(cleaned), 1) >= 0.6


def is_heading_1_text(text: str) -> bool:
    text = normalize_text(text)
    if not text:
        return False
    if text in {TEXT["intro"], TEXT["conclusion"]}:
        return True
    if len(text) > 60:
        return False
    if re.match(r"^(0|[1-9]\d*)\s{0,3}\S+", text) and not re.match(r"^\d+\.\d+", text):
        return True
    return False


def is_heading_2_text(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+", normalize_text(text)))


def _style_name(paragraph) -> str:
    return (paragraph.style.name or "").strip().lower() if paragraph.style is not None and paragraph.style.name else ""


def _is_short_heading_candidate(text: str) -> bool:
    text = normalize_text(text)
    if not text or len(text) > 40:
        return False
    if text.startswith(SKIP_PREFIXES):
        return False
    if is_reference_entry_text(text) or is_caption_like_text(text):
        return False
    if text.startswith((TEXT["references"], TEXT["author_bio"], TEXT["received"], TEXT["fund"])):
        return False
    if text.endswith(("。", "；", ";", "！", "？", "：", ":")):
        return False
    return True


def is_heading_1_paragraph(paragraph) -> bool:
    text = normalize_text(paragraph.text)
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return False
    if text.lower().startswith("doi:") or text.startswith((TEXT["abstract"], TEXT["keywords"], TEXT["keywords_alt"], TEXT["references"], TEXT["author_bio"], TEXT["received"], TEXT["fund"])):
        return False
    style_name = _style_name(paragraph)
    if style_name in {"heading 2", "标题 2"}:
        return False
    if is_heading_1_text(text):
        return True
    if style_name in {"heading 1", "标题 1"}:
        return True
    if not _is_short_heading_candidate(text):
        return False
    if "：" in text or ":" in text:
        return False
    snap = paragraph_snapshot(paragraph)
    if snap.get("line_spacing_mode") == "multiple" and snap.get("line_spacing_value") == 2.0:
        return True
    return False


def is_heading_2_paragraph(paragraph) -> bool:
    text = normalize_text(paragraph.text)
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return False
    if text.lower().startswith("doi:") or text.startswith((TEXT["abstract"], TEXT["keywords"], TEXT["keywords_alt"], TEXT["references"], TEXT["author_bio"], TEXT["received"], TEXT["fund"])):
        return False
    if is_heading_2_text(text):
        return True
    style_name = _style_name(paragraph)
    if style_name in {"heading 2", "标题 2"}:
        return True
    if not _is_short_heading_candidate(text):
        return False
    snap = paragraph_snapshot(paragraph)
    if snap.get("font_east_asia") == "黑体" or snap.get("font_ascii") == "黑体":
        return True
    return False


def is_reference_entry_text(text: str) -> bool:
    text = normalize_text(text)
    return bool(re.match(r"^\[\d+\]", text)) or bool(re.match(r"^[A-Z][A-Z\s,.-]{8,}", text))


def is_caption_like_text(text: str) -> bool:
    text = normalize_text(text)
    return bool(re.match(rf"^({TEXT['figure']}|{TEXT['table']})\s*\d+", text)) or bool(
        re.match(r"^(Figure|Table|Fig\.)\s*\d+", text, re.I)
    )


def is_instructional_text(text: str) -> bool:
    text = normalize_text(text)
    if not text:
        return False
    formatting_terms = [
        "字体",
        "字号",
        "行距",
        "段后",
        "段前",
        "缩进",
        "居中",
        "排列",
        "半角",
        "双栏",
        "栏间距",
        "页眉",
        "页脚",
        "黑体",
        "宋体",
        "楷体",
        "Times New Roman",
        "Arial",
        "MathType",
        "Visio",
        "公式序号",
        "正体",
        "斜体",
        "右对齐",
        "三线表",
        "示例如下",
        "著录要求",
        "著录格式",
        "要全部指明",
        "小五",
        "五号",
        "小四",
        "小二",
        "论文正文",
        "参考文献格式",
    ]
    hits = sum(1 for term in formatting_terms if term.lower() in text.lower())
    if hits >= 2:
        return True
    if "对应" in text and "字体" in text:
        return True
    if "格式" in text and "著录" in text:
        return True
    return False


def _line_spacing_info_from_ppr(p_pr) -> Dict[str, Any]:
    if p_pr is None or p_pr.spacing is None:
        return {"line_spacing_mode": None, "line_spacing_value": None}
    raw_line = p_pr.spacing.get(qn("w:line"))
    if raw_line is None:
        return {"line_spacing_mode": None, "line_spacing_value": None}
    line = int(raw_line)
    line_rule = p_pr.spacing.get(qn("w:lineRule")) or "auto"
    if line_rule == "auto":
        return {"line_spacing_mode": "multiple", "line_spacing_value": round(line / 240.0, 2)}
    if line_rule == "exact":
        return {"line_spacing_mode": "exact", "line_spacing_value": round(line / 20.0, 2)}
    if line_rule == "atLeast":
        return {"line_spacing_mode": "at_least", "line_spacing_value": round(line / 20.0, 2)}
    return {"line_spacing_mode": line_rule, "line_spacing_value": round(line / 20.0, 2)}


def _line_spacing_info(paragraph) -> Dict[str, Any]:
    direct = _line_spacing_info_from_ppr(paragraph._p.pPr)
    if direct["line_spacing_value"] is not None:
        return direct
    for style in _style_chain(paragraph.style):
        p_pr = getattr(style.element, "pPr", None)
        info = _line_spacing_info_from_ppr(p_pr)
        if info["line_spacing_value"] is not None:
            return info
    return direct


def _indent_info_from_ppr(p_pr) -> Dict[str, Any]:
    ind = p_pr.ind if p_pr is not None else None
    left_indent = right_indent = first_line = hanging = None
    if ind is not None:
        raw_left = ind.get(qn("w:left"))
        raw_right = ind.get(qn("w:right"))
        raw_first = ind.get(qn("w:firstLine"))
        raw_hanging = ind.get(qn("w:hanging"))
        if raw_left is not None:
            left_indent = round(int(raw_left) / 20.0, 2)
        if raw_right is not None:
            right_indent = round(int(raw_right) / 20.0, 2)
        if raw_first is not None:
            first_line = round(int(raw_first) / 20.0, 2)
        if raw_hanging is not None:
            hanging = round(int(raw_hanging) / 20.0, 2)
    return {
        "left_indent_pt": left_indent,
        "right_indent_pt": right_indent,
        "first_line_indent_pt": first_line,
        "hanging_indent_pt": hanging,
    }


def _indent_info(paragraph) -> Dict[str, Any]:
    direct = _indent_info_from_ppr(paragraph._p.pPr)
    if any(direct.values()):
        return direct
    for style in _style_chain(paragraph.style):
        p_pr = getattr(style.element, "pPr", None)
        info = _indent_info_from_ppr(p_pr)
        if any(info.values()):
            return info
    return direct


def _style_length(style, attr: str) -> Optional[float]:
    for candidate in _style_chain(style):
        fmt = getattr(candidate, "paragraph_format", None)
        if fmt is None:
            continue
        value = getattr(fmt, attr, None)
        if value is not None:
            return _fmt_length(value)
    return None


def _bool_from_on_off(element) -> Optional[bool]:
    if element is None:
        return None
    val = element.get(qn("w:val"))
    if val is None:
        return True
    return val not in {"0", "false", "False", "off"}


def _font_info_from_rpr(r_pr) -> Dict[str, Any]:
    if r_pr is None:
        return {}
    r_fonts = getattr(r_pr, "rFonts", None)
    size = None
    sz = r_pr.find(qn("w:sz"))
    if sz is not None and sz.get(qn("w:val")):
        size = round(int(sz.get(qn("w:val"))) / 2.0, 2)
    elif r_pr.find(qn("w:szCs")) is not None and r_pr.find(qn("w:szCs")).get(qn("w:val")):
        size = round(int(r_pr.find(qn("w:szCs")).get(qn("w:val"))) / 2.0, 2)
    return {
        "font_east_asia": r_fonts.get(qn("w:eastAsia")) if r_fonts is not None else None,
        "font_ascii": r_fonts.get(qn("w:ascii")) if r_fonts is not None else None,
        "font_hansi": r_fonts.get(qn("w:hAnsi")) if r_fonts is not None else None,
        "font_cs": r_fonts.get(qn("w:cs")) if r_fonts is not None else None,
        "font_size_pt": size,
        "bold": _bool_from_on_off(r_pr.find(qn("w:b"))),
        "italic": _bool_from_on_off(r_pr.find(qn("w:i"))),
    }


def _style_font_info(paragraph) -> Dict[str, Any]:
    merged = {
        "font_east_asia": None,
        "font_ascii": None,
        "font_hansi": None,
        "font_cs": None,
        "font_size_pt": None,
        "bold": None,
        "italic": None,
    }
    seen_ids = set()
    styles_to_check = list(_style_chain(paragraph.style))
    try:
        normal_style = paragraph.part.document.styles["Normal"]
    except Exception:
        normal_style = None
    if normal_style is not None:
        styles_to_check.append(normal_style)

    for style in styles_to_check:
        style_id = getattr(style, "style_id", id(style))
        if style_id in seen_ids:
            continue
        seen_ids.add(style_id)
        info = _font_info_from_rpr(getattr(style.element, "rPr", None))
        for key, value in info.items():
            if merged.get(key) is None and value is not None:
                merged[key] = value
    return merged


def _docdefaults_font_info(paragraph) -> Dict[str, Any]:
    merged = {
        "font_east_asia": None,
        "font_ascii": None,
        "font_hansi": None,
        "font_cs": None,
        "font_size_pt": None,
        "bold": None,
        "italic": None,
    }
    try:
        styles_element = paragraph.part.document.styles.element
    except Exception:
        return merged

    doc_defaults = styles_element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        return merged
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        return merged
    r_pr = rpr_default.find(qn("w:rPr"))
    info = _font_info_from_rpr(r_pr)
    for key, value in info.items():
        if value is not None:
            merged[key] = value
    return merged


def _font_info_from_runs(paragraph) -> Dict[str, Any]:
    east_asia: List[str] = []
    ascii_names: List[str] = []
    hansi_names: List[str] = []
    sizes: List[float] = []
    bolds: List[bool] = []
    italics: List[bool] = []
    for run in paragraph.runs:
        text = normalize_text(run.text)
        if not text:
            continue
        r_pr = run._element.rPr
        if r_pr is not None and r_pr.rFonts is not None:
            east = r_pr.rFonts.get(qn("w:eastAsia"))
            ascii_font = r_pr.rFonts.get(qn("w:ascii"))
            hansi = r_pr.rFonts.get(qn("w:hAnsi"))
            if east:
                east_asia.append(east)
            if ascii_font:
                ascii_names.append(ascii_font)
            if hansi:
                hansi_names.append(hansi)
        if run.font.name:
            if contains_cjk(text):
                east_asia.append(run.font.name)
            else:
                ascii_names.append(run.font.name)
        if run.font.size:
            sizes.append(round(run.font.size.pt, 2))
        if run.bold is not None:
            bolds.append(bool(run.bold))
        if run.italic is not None:
            italics.append(bool(run.italic))
    info = {
        "font_east_asia": Counter(east_asia).most_common(1)[0][0] if east_asia else None,
        "font_ascii": Counter(ascii_names).most_common(1)[0][0] if ascii_names else None,
        "font_hansi": Counter(hansi_names).most_common(1)[0][0] if hansi_names else None,
        "font_size_pt": round(median(sizes), 2) if sizes else None,
        "bold": Counter(bolds).most_common(1)[0][0] if bolds else None,
        "italic": Counter(italics).most_common(1)[0][0] if italics else None,
    }
    for fallback in (_style_font_info(paragraph), _docdefaults_font_info(paragraph)):
        for key in ("font_east_asia", "font_ascii", "font_hansi", "font_size_pt", "bold", "italic"):
            if info.get(key) is None and fallback.get(key) is not None:
                info[key] = fallback[key]
    return info


def paragraph_snapshot(paragraph) -> Dict[str, Any]:
    fmt = paragraph.paragraph_format
    snap = {
        "style_name": paragraph.style.name if paragraph.style else None,
        "alignment": ALIGNMENT_MAP.get(paragraph.alignment, None),
        "space_before_pt": _fmt_length(fmt.space_before) if fmt.space_before is not None else _style_length(paragraph.style, "space_before"),
        "space_after_pt": _fmt_length(fmt.space_after) if fmt.space_after is not None else _style_length(paragraph.style, "space_after"),
    }
    snap.update(_line_spacing_info(paragraph))
    snap.update(_indent_info(paragraph))
    snap.update(_font_info_from_runs(paragraph))
    return snap


def _aggregate_value(values: List[Any]) -> Any:
    filtered = [value for value in values if value is not None]
    if not filtered:
        return None
    first = filtered[0]
    if isinstance(first, bool):
        counter = Counter(filtered)
        value, count = counter.most_common(1)[0]
        return value if count / len(filtered) >= 0.75 else None
    if isinstance(first, (int, float)):
        return round(median(filtered), 2)
    counter = Counter(filtered)
    value, count = counter.most_common(1)[0]
    return value if count / len(filtered) >= 0.5 else value


def aggregate_snapshots(paragraphs: Iterable) -> Dict[str, Any]:
    paragraphs = list(paragraphs)
    if not paragraphs:
        return {}
    snaps = [paragraph_snapshot(paragraph) for paragraph in paragraphs]
    keys = set().union(*(snap.keys() for snap in snaps))
    merged = {key: _aggregate_value([snap.get(key) for snap in snaps]) for key in keys}
    merged["sample_count"] = len(paragraphs)
    merged["sample_texts"] = [normalize_text(paragraph.text)[:120] for paragraph in paragraphs[:5]]
    return merged


def section_layouts(doc: Document) -> List[Dict[str, Any]]:
    layouts: List[Dict[str, Any]] = []
    break_indices: List[int] = []
    for index, paragraph in enumerate(doc.paragraphs):
        p_pr = paragraph._p.pPr
        has_sect = p_pr is not None and p_pr.find(qn("w:sectPr")) is not None
        if has_sect:
            break_indices.append(index)

    section_starts = [0]
    section_starts.extend(index + 1 for index in break_indices)
    while len(section_starts) > len(doc.sections):
        section_starts.pop()

    for index, section in enumerate(doc.sections):
        sect_pr = section._sectPr
        cols = sect_pr.find(qn("w:cols"))
        num = int(cols.get(qn("w:num"))) if cols is not None and cols.get(qn("w:num")) else 1
        space_twips = int(cols.get(qn("w:space"))) if cols is not None and cols.get(qn("w:space")) else 0
        start_index = section_starts[index] if index < len(section_starts) else None
        end_index = (section_starts[index + 1] - 1) if index + 1 < len(section_starts) else len(doc.paragraphs) - 1
        layouts.append(
            {
                "section_index": index,
                "start_paragraph_index": start_index,
                "end_paragraph_index": end_index,
                "columns_num": num,
                "columns_space_pt": round(space_twips / 20.0, 2),
                "page_width": _fmt_length(section.page_width),
                "page_height": _fmt_length(section.page_height),
                "top_margin": _fmt_length(section.top_margin),
                "bottom_margin": _fmt_length(section.bottom_margin),
                "left_margin": _fmt_length(section.left_margin),
                "right_margin": _fmt_length(section.right_margin),
                "header_distance": _fmt_length(section.header_distance),
                "footer_distance": _fmt_length(section.footer_distance),
            }
        )
    return layouts


def filter_instructional_paragraphs(paragraphs: Iterable, role: str) -> List:
    paragraphs = list(paragraphs)
    if not paragraphs:
        return []

    prefer_filtered_roles = {
        "heading_1",
        "heading_2",
        "body",
        "caption_figure",
        "caption_table",
        "references_title",
        "reference_entry",
        "fund",
        "author_bio_title",
        "author_bio_entry",
    }
    if role not in prefer_filtered_roles:
        return paragraphs

    filtered = [paragraph for paragraph in paragraphs if not is_instructional_text(paragraph.text)]
    return filtered or paragraphs


def collect_non_empty_paragraphs(doc: Document) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = normalize_text(paragraph.text)
        if not text:
            continue
        items.append(
            {
                "index": index,
                "text": text,
                "style_name": paragraph.style.name if paragraph.style else None,
                "alignment": ALIGNMENT_MAP.get(paragraph.alignment, None),
                "paragraph": paragraph,
            }
        )
    return items


def detect_blocks(doc: Document) -> Dict[str, List]:
    items = collect_non_empty_paragraphs(doc)
    blocks: Dict[str, List] = {key: [] for key in ROLE_LABELS}

    def add(role: str, paragraph) -> None:
        if paragraph is None:
            return
        if paragraph not in blocks[role]:
            blocks[role].append(paragraph)

    if items:
        for item in items[:12]:
            text = item["text"]
            if text.lower().startswith("doi:"):
                continue
            if text.startswith((TEXT["abstract"], TEXT["keywords"], TEXT["keywords_alt"])):
                continue
            add("title", item["paragraph"])
            break

    title_index = next((item["index"] for item in items if blocks["title"] and item["paragraph"] is blocks["title"][0]), None)
    if title_index is not None:
        following = [item for item in items if item["index"] > title_index]
        centered = [item for item in following if item["alignment"] == "center"]
        if centered:
            add("authors", centered[0]["paragraph"])
        if len(centered) > 1:
            add("affiliations", centered[1]["paragraph"])

    ref_title_index = None
    for item in items:
        lower_text = item["text"].lower()
        paragraph = item["paragraph"]
        if item["text"].startswith(TEXT["abstract"]):
            add("abstract_cn", paragraph)
        elif item["text"].startswith((TEXT["keywords"], TEXT["keywords_alt"])):
            add("keywords_cn", paragraph)
        elif lower_text.startswith("abstract"):
            add("abstract_en", paragraph)
        elif lower_text.startswith("key words") or lower_text.startswith("keywords"):
            add("keywords_en", paragraph)
        elif item["text"].startswith(TEXT["fund"]):
            add("fund", paragraph)
        elif item["text"].startswith(TEXT["author_bio"]):
            add("author_bio_title", paragraph)
        elif item["text"].startswith(TEXT["received"]):
            add("footer", paragraph)
        elif item["text"].startswith(TEXT["references"]):
            add("references_title", paragraph)
            ref_title_index = item["index"]
        elif (ref_title_index is None or item["index"] < ref_title_index) and is_heading_1_paragraph(paragraph):
            add("heading_1", paragraph)
        elif (ref_title_index is None or item["index"] < ref_title_index) and is_heading_2_paragraph(paragraph):
            add("heading_2", paragraph)

    keyword_index = next((item["index"] for item in items if blocks["keywords_cn"] and item["paragraph"] is blocks["keywords_cn"][0]), None)
    if keyword_index is not None:
        for item in items:
            if item["index"] <= keyword_index:
                continue
            if is_ascii_heavy(item["text"]) and item["alignment"] == "center":
                add("title_en", item["paragraph"])
                break
    title_en_index = next((item["index"] for item in items if blocks["title_en"] and item["paragraph"] is blocks["title_en"][0]), None)
    if title_en_index is not None:
        following_centered = [item for item in items if item["index"] > title_en_index and item["alignment"] == "center"]
        if following_centered:
            add("authors_en", following_centered[0]["paragraph"])
        if len(following_centered) > 1:
            add("affiliations_en", following_centered[1]["paragraph"])

    ref_title_index = next((item["index"] for item in items if blocks["references_title"] and item["paragraph"] is blocks["references_title"][0]), ref_title_index)
    if ref_title_index is not None:
        for item in items:
            if item["index"] <= ref_title_index:
                continue
            if item["text"].startswith((TEXT["received"], TEXT["author_bio"])):
                break
            if is_reference_entry_text(item["text"]):
                add("reference_entry", item["paragraph"])

    author_bio_title_index = next((item["index"] for item in items if blocks["author_bio_title"] and item["paragraph"] is blocks["author_bio_title"][0]), None)
    if author_bio_title_index is not None:
        for item in items:
            if item["index"] <= author_bio_title_index:
                continue
            if re.match(r"^[^()\uFF08\uFF09]{2,20}[\uFF08(]", item["text"]):
                add("author_bio_entry", item["paragraph"])

    for item in items:
        text = item["text"]
        paragraph = item["paragraph"]
        if TEXT["figure_caption"] in text or re.match(rf"^{TEXT['figure']}\s*\d+", text) or re.match(r"^(Figure|Fig\.)\s*\d+", text, re.I):
            add("caption_figure", paragraph)
        elif TEXT["table_caption"] in text or re.match(rf"^{TEXT['table']}\s*\d+", text) or re.match(r"^Table\s*\d+", text, re.I):
            add("caption_table", paragraph)

    body_candidates = []
    for item in items:
        text = item["text"]
        lower_text = text.lower()
        if ref_title_index is not None and item["index"] >= ref_title_index:
            continue
        if len(text) < 18:
            continue
        if lower_text.startswith(SKIP_PREFIXES):
            continue
        if item["alignment"] == "center":
            continue
        if is_heading_1_paragraph(item["paragraph"]) or is_heading_2_paragraph(item["paragraph"]):
            continue
        if is_reference_entry_text(text) or is_caption_like_text(text):
            continue
        if TEXT["figure_caption"] in text or TEXT["table_caption"] in text:
            continue
        body_candidates.append(item["paragraph"])
    blocks["body"] = body_candidates[:12]

    if doc.sections:
        header_para = next((p for p in doc.sections[0].header.paragraphs if normalize_text(p.text)), None)
        footer_para = next((p for p in doc.sections[0].footer.paragraphs if normalize_text(p.text)), None)
        add("header", header_para)
        if footer_para is not None:
            add("footer", footer_para)

    # These roles should represent a single anchor block rather than a list of
    # all similarly prefixed paragraphs.
    for role in ("references_title", "author_bio_title", "fund", "header", "footer"):
        if len(blocks[role]) > 1:
            blocks[role] = blocks[role][:1]

    return {role: paragraphs for role, paragraphs in blocks.items() if paragraphs}


def build_block_rule(paragraphs: List) -> Dict[str, Any]:
    return aggregate_snapshots(paragraphs)
