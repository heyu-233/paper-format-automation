from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from docx_rule_utils import detect_blocks, filter_instructional_paragraphs, is_heading_1_paragraph, is_heading_2_paragraph

TEXT = {
    "abstract": "\u6458\u8981",
    "keywords": "\u5173\u952e\u8bcd",
    "keywords_alt": "\u5173\u952e\u5b57",
    "references": "\u53c2\u8003\u6587\u732e",
    "author_bio": "\u4f5c\u8005\u7b80\u4ecb",
    "figure_caption": "\u56fe\u9898",
    "table_caption": "\u8868\u9898",
}

TEST_FRONT_MATTER = {
    "authors": "\u738b\u660e1\uff0c\u674e\u534e2",
    "affiliations": "(1.\u793a\u4f8b\u5355\u4f4d\uff0c\u5317\u4eac 100000\uff1b2.\u793a\u4f8b\u5355\u4f4d\uff0c\u8d35\u9633 550000)",
    "title_en": "Health Protection System Innovation for Frontline Power Workers",
    "authors_en": "WANG Ming1, LI Hua2",
    "affiliations_en": "(1. Example Institute, Beijing 100000, China; 2. Example Institute, Guiyang 550000, China)",
    "abstract_en": "Abstract: Placeholder abstract for front-matter layout testing.",
    "keywords_en": "Key words: power workers; health protection; system innovation.",
}

TEST_CAPTIONS = {
    "figure_cn": "\u56fe{index}",
    "figure_en": "Figure {index}.",
    "table_cn": "\u8868{index}",
    "table_en": "Table {index}.",
}

TEST_OPTIONAL_BLOCKS = {
    "fund": "\u57fa\u91d1\u9879\u76ee\uff1a\u57fa\u91d1\u4fe1\u606f\u5f85\u8865\u5145",
    "footer": "\u6536\u7a3f\u65e5\u671f\uff1a2026-04-16\uff1b    \u4fee\u56de\u65e5\u671f\uff1a2026-04-16",
    "author_bio_title": "\u4f5c\u8005\u7b80\u4ecb\uff1a",
    "author_bio_entry_1": "\u738b\u660e(1990-)\uff0c\u7537\uff0c\u901a\u4fe1\u4f5c\u8005\uff0c\u535a\u58eb\uff0c\u7814\u7a76\u65b9\u5411\u4e3a\u7535\u529b\u4f5c\u4e1a\u4eba\u5458\u5065\u5eb7\u4fdd\u969c\u4e0e\u98ce\u9669\u9884\u8b66\uff1bE-mail: wangming@example.com",
    "author_bio_entry_2": "\u674e\u534e(1992-)\uff0c\u5973\uff0c\u7855\u58eb\uff0c\u7814\u7a76\u65b9\u5411\u4e3a\u7535\u529b\u4f5c\u4e1a\u5065\u5eb7\u6570\u636e\u6cbb\u7406\u4e0e\u667a\u80fd\u9884\u8b66\u3002E-mail: lihua@example.com",
    "author_bio_entry_3": "\u5f20\u5f3a(1988-)\uff0c\u7537\uff0c\u7855\u58eb\uff0c\u7814\u7a76\u65b9\u5411\u4e3a\u7535\u529b\u4f5c\u4e1a\u98ce\u9669\u8bc6\u522b\u4e0e\u5b89\u5168\u7ba1\u7406\u3002E-mail: zhangqiang@example.com",
}

ALIGNMENT_REVERSE = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

UNIFORM_BLOCK_ROLES = {
    "title",
    "authors",
    "affiliations",
    "abstract_cn",
    "keywords_cn",
    "title_en",
    "authors_en",
    "affiliations_en",
    "abstract_en",
    "keywords_en",
    "heading_1",
    "heading_2",
    "fund",
    "caption_figure",
    "caption_table",
    "references_title",
    "reference_entry",
    "author_bio_title",
    "author_bio_entry",
    "header",
    "footer",
}


def _apply_run_style(run, rule: Dict[str, Any]) -> None:
    east_asia = rule.get("font_east_asia")
    ascii_font = rule.get("font_ascii")
    hansi_font = rule.get("font_hansi")
    cs_font = rule.get("font_cs")
    if east_asia or ascii_font or hansi_font or cs_font:
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        if east_asia:
            r_fonts.set(qn("w:eastAsia"), east_asia)
        if ascii_font:
            r_fonts.set(qn("w:ascii"), ascii_font)
        if hansi_font:
            r_fonts.set(qn("w:hAnsi"), hansi_font)
        elif ascii_font:
            # Only fall back when the template did not provide a separate hAnsi family.
            r_fonts.set(qn("w:hAnsi"), ascii_font)
        if cs_font:
            r_fonts.set(qn("w:cs"), cs_font)
    if rule.get("font_size_pt") is not None:
        run.font.size = Pt(rule["font_size_pt"])
    if rule.get("bold") is not None:
        run.bold = bool(rule["bold"])
    if rule.get("italic") is not None:
        run.italic = bool(rule["italic"])


def _apply_paragraph_style(paragraph, rule: Dict[str, Any]) -> None:
    fmt = paragraph.paragraph_format
    alignment = rule.get("alignment")
    if alignment in ALIGNMENT_REVERSE:
        paragraph.alignment = ALIGNMENT_REVERSE[alignment]
    line_mode = rule.get("line_spacing_mode")
    line_value = rule.get("line_spacing_value")
    if line_mode == "multiple" and line_value is not None:
        fmt.line_spacing = line_value
    elif line_mode == "exact" and line_value is not None:
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(line_value)
    elif line_mode == "at_least" and line_value is not None:
        fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        fmt.line_spacing = Pt(line_value)
    if rule.get("space_before_pt") is not None:
        fmt.space_before = Pt(rule["space_before_pt"])
    if rule.get("space_after_pt") is not None:
        fmt.space_after = Pt(rule["space_after_pt"])
    if rule.get("left_indent_pt") is not None:
        fmt.left_indent = Pt(rule["left_indent_pt"])
    if rule.get("right_indent_pt") is not None:
        fmt.right_indent = Pt(rule["right_indent_pt"])
    if rule.get("first_line_indent_pt") is not None:
        fmt.first_line_indent = Pt(rule["first_line_indent_pt"])
    elif rule.get("hanging_indent_pt") is not None:
        fmt.first_line_indent = Pt(-rule["hanging_indent_pt"])
    for run in paragraph.runs:
        if (run.text or "").strip():
            _apply_run_style(run, rule)


def _text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\xa0", " ")).strip()


def _is_heading_1(text: str) -> bool:
    text = _text(text)
    if text in {"引言", "结论"}:
        return True
    return bool(re.match(r"^(0|[1-9]\d*)\s{0,3}\S+", text)) and not bool(re.match(r"^\d+\.\d+", text))


def _is_heading_2(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+", _text(text)))


def _is_reference(text: str) -> bool:
    text = _text(text)
    return bool(re.match(r"^\[\d+\]", text)) or bool(re.match(r"^[A-Z][A-Z\s,.-]{8,}", text))


def _is_centered_paragraph(paragraph) -> bool:
    return paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER


def _find_first_nonempty_paragraph(doc: Document):
    for paragraph in doc.paragraphs:
        if _text(paragraph.text):
            return paragraph
    return None


def _find_paragraph_starting(doc: Document, prefixes):
    if isinstance(prefixes, str):
        prefixes = (prefixes,)
    lowered = tuple(prefix.lower() for prefix in prefixes)
    for paragraph in doc.paragraphs:
        text = _text(paragraph.text)
        if text and text.lower().startswith(lowered):
            return paragraph
    return None


def _clear_story(story) -> None:
    element = story._element
    for child in list(element):
        element.remove(child)


def _copy_run_properties(src_run, dst_run) -> None:
    src_rpr = src_run._element.rPr
    if src_rpr is None:
        return
    dst_r = dst_run._element
    dst_rpr = dst_r.rPr
    if dst_rpr is not None:
        dst_r.remove(dst_rpr)
    dst_r.insert(0, deepcopy(src_rpr))


def _clear_run_properties(run) -> None:
    r_pr = run._element.rPr
    if r_pr is not None:
        run._element.remove(r_pr)


def _style_name(style) -> str | None:
    return style.name if style is not None and style.name else None


def _sync_template_style(doc: Document, template_doc: Document, style_name: str | None) -> None:
    if not style_name:
        return
    try:
        template_style = template_doc.styles[style_name]
    except Exception:
        return

    styles_root = doc.styles.element
    template_style_element = deepcopy(template_style.element)
    template_style_id = template_style.style_id

    for existing in list(styles_root.findall(qn("w:style"))):
        existing_name_el = existing.find(qn("w:name"))
        existing_name = existing_name_el.get(qn("w:val")) if existing_name_el is not None else None
        if existing.get(qn("w:styleId")) == template_style_id or existing_name == style_name:
            styles_root.remove(existing)
    styles_root.append(template_style_element)


def _sync_template_styles(doc: Document, template_doc: Document, template_blocks: Dict[str, Any]) -> None:
    if template_doc is None:
        return
    style_names = {"Normal"}
    for paragraph in template_blocks.values():
        style_name = _style_name(getattr(paragraph, "style", None))
        if style_name:
            style_names.add(style_name)
    for section in template_doc.sections:
        for story in (section.header, section.first_page_header, section.footer, section.first_page_footer):
            for paragraph in story.paragraphs:
                style_name = _style_name(getattr(paragraph, "style", None))
                if style_name:
                    style_names.add(style_name)
    for style_name in style_names:
        _sync_template_style(doc, template_doc, style_name)


def _rewrite_story_from_template(src_story, dst_story) -> None:
    # Force python-docx to materialize a valid target part before editing it.
    _ = dst_story.paragraphs
    _clear_story(dst_story)
    for src_paragraph in src_story.paragraphs:
        dst_paragraph = dst_story.add_paragraph()
        _copy_paragraph_properties(src_paragraph, dst_paragraph)
        if src_paragraph.style is not None:
            try:
                dst_paragraph.style = src_paragraph.style.name
            except Exception:
                pass
        for src_run in src_paragraph.runs:
            dst_run = dst_paragraph.add_run(src_run.text)
            _copy_run_properties(src_run, dst_run)
        if not src_paragraph.runs and src_paragraph.text:
            dst_paragraph.add_run(src_paragraph.text)


def _copy_paragraph_properties(src_paragraph, dst_paragraph) -> None:
    if src_paragraph is None or dst_paragraph is None:
        return
    dst_p = dst_paragraph._p
    existing = dst_p.pPr
    preserved_sectpr = None
    if existing is not None:
        sect_pr = existing.find(qn("w:sectPr"))
        if sect_pr is not None:
            preserved_sectpr = deepcopy(sect_pr)
    if existing is not None:
        dst_p.remove(existing)
    src_ppr = src_paragraph._p.pPr
    if src_ppr is not None:
        copied = deepcopy(src_ppr)
        for sect_pr in copied.findall(qn("w:sectPr")):
            copied.remove(sect_pr)
        if preserved_sectpr is not None:
            copied.append(preserved_sectpr)
        dst_p.insert(0, copied)
    elif preserved_sectpr is not None:
        dst_p.insert(0, preserved_sectpr)


def _apply_template_block(paragraph, rule: Dict[str, Any] | None, template_paragraph) -> None:
    role = rule.get("role") if rule else None
    if template_paragraph is not None:
        try:
            if template_paragraph.style is not None and template_paragraph.style.name:
                paragraph.style = template_paragraph.style.name
        except Exception:
            pass
        _copy_paragraph_properties(template_paragraph, paragraph)
    if role in UNIFORM_BLOCK_ROLES:
        for run in paragraph.runs:
            if (run.text or "").strip():
                _clear_run_properties(run)
    if rule:
        _apply_paragraph_style(paragraph, rule)


def _insert_front_block(before_paragraph, text: str, rule: Dict[str, Any] | None):
    paragraph = before_paragraph.insert_paragraph_before(text)
    if rule:
        _apply_paragraph_style(paragraph, rule)
    return paragraph


def _insert_front_blocks(before_paragraph, items):
    for text, rule in items:
        _insert_front_block(before_paragraph, text, rule)


def _insert_block_after(after_paragraph, text: str, rule: Dict[str, Any] | None, template_paragraph=None):
    paragraph = _new_paragraph_after(after_paragraph, text)
    _apply_template_block(paragraph, rule, template_paragraph)
    return paragraph


def _paragraph_index(doc: Document, target) -> int | None:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph._p is target._p:
            return index
    return None


def _new_paragraph_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def _new_paragraph_before_table(table, text: str):
    new_p = OxmlElement("w:p")
    table._tbl.addprevious(new_p)
    new_para = Paragraph(new_p, table._parent)
    if text:
        new_para.add_run(text)
    return new_para


def _collect_existing_caption_numbers(paragraphs, pattern: str) -> set[int]:
    seen = set()
    for paragraph in paragraphs:
        text = _text(paragraph.text)
        match = re.match(pattern, text, re.I)
        if match:
            seen.add(int(match.group(1)))
    return seen


def _find_nearby_number(paragraphs, index: int, pattern: str) -> int | None:
    for offset in (-3, -2, -1, 1, 2, 3):
        check = index + offset
        if 0 <= check < len(paragraphs):
            text = _text(paragraphs[check].text)
            match = re.search(pattern, text, re.I)
            if match:
                return int(match.group(1))
    return None


def _next_unused(preferred: int | None, used: set[int]) -> int:
    if preferred is not None and preferred not in used:
        used.add(preferred)
        return preferred
    candidate = 1
    while candidate in used:
        candidate += 1
    used.add(candidate)
    return candidate


def _ensure_front_matter(doc: Document, blocks: Dict[str, Dict[str, Any]]) -> None:
    title_paragraph = _find_first_nonempty_paragraph(doc)
    abstract_cn = _find_paragraph_starting(doc, TEXT["abstract"])
    keywords_cn = _find_paragraph_starting(doc, (TEXT["keywords"], TEXT["keywords_alt"]))
    if title_paragraph is None or abstract_cn is None or keywords_cn is None:
        return

    title_index = _paragraph_index(doc, title_paragraph)
    abstract_index = _paragraph_index(doc, abstract_cn)
    if title_index is None or abstract_index is None:
        return
    front_slice = doc.paragraphs[title_index + 1 : abstract_index]
    centered_cn = [paragraph for paragraph in front_slice if _is_centered_paragraph(paragraph)]
    if not centered_cn:
        _insert_front_blocks(
            abstract_cn,
            [
                (TEST_FRONT_MATTER["authors"], blocks.get("authors")),
                (TEST_FRONT_MATTER["affiliations"], blocks.get("affiliations")),
            ],
        )

    english_title = _find_paragraph_starting(doc, "english title")
    abstract_en = _find_paragraph_starting(doc, "abstract")
    keywords_en = _find_paragraph_starting(doc, ("key words", "keywords"))
    intro_anchor = _find_paragraph_starting(doc, "引言")
    english_anchor = abstract_en or intro_anchor or keywords_cn
    if english_title is None:
        _insert_front_blocks(
            english_anchor,
            [
                (TEST_FRONT_MATTER["title_en"], blocks.get("title_en")),
                (TEST_FRONT_MATTER["authors_en"], blocks.get("authors_en")),
                (TEST_FRONT_MATTER["affiliations_en"], blocks.get("affiliations_en")),
            ],
        )
    if abstract_en is None:
        abstract_anchor = _find_paragraph_starting(doc, "引言") or keywords_cn
        _insert_front_blocks(
            abstract_anchor,
            [
                (TEST_FRONT_MATTER["abstract_en"], blocks.get("abstract_en")),
                (TEST_FRONT_MATTER["keywords_en"], blocks.get("keywords_en")),
            ],
        )


def _has_nonempty_neighbor(paragraphs, index: int, patterns) -> bool:
    for offset in (-2, -1, 1, 2):
        check = index + offset
        if 0 <= check < len(paragraphs):
            text = _text(paragraphs[check].text)
            if text and any(re.match(pattern, text, re.I) for pattern in patterns):
                return True
    return False


def _ensure_caption_placeholders(doc: Document, blocks: Dict[str, Dict[str, Any]]) -> None:
    paragraphs = doc.paragraphs
    figure_patterns = [r"^图\s*\d+", r"^(Figure|Fig\.)\s*\d+"]
    table_patterns = [r"^表\s*\d+", r"^Table\s*\d+"]
    used_figure_numbers = _collect_existing_caption_numbers(paragraphs, r"^(?:图|Figure|Fig\.)\s*(\d+)")
    used_table_numbers = _collect_existing_caption_numbers(paragraphs, r"^(?:表|Table)\s*(\d+)")

    for paragraph in list(paragraphs):
        has_drawing = bool(paragraph._p.findall(".//" + qn("w:drawing")))
        if not has_drawing:
            continue
        current_index = _paragraph_index(doc, paragraph)
        if current_index is None:
            continue
        if _has_nonempty_neighbor(doc.paragraphs, current_index, figure_patterns):
            continue
        figure_index = _next_unused(_find_nearby_number(doc.paragraphs, current_index, r"图\s*(\d+)"), used_figure_numbers)
        en_para = _new_paragraph_after(paragraph, TEST_CAPTIONS["figure_en"].format(index=figure_index))
        cn_para = _new_paragraph_after(paragraph, TEST_CAPTIONS["figure_cn"].format(index=figure_index))
        paragraph.paragraph_format.keep_with_next = True
        cn_para.paragraph_format.keep_with_next = True
        en_para.paragraph_format.keep_with_next = True
        en_para.paragraph_format.keep_together = True
        if blocks.get("caption_figure"):
            _apply_paragraph_style(cn_para, blocks["caption_figure"])
            _apply_paragraph_style(en_para, blocks["caption_figure"])

    for table in doc.tables:
        prev = table._tbl.getprevious()
        next_ = table._tbl.getnext()
        near_text = []
        for node in (prev, next_):
            if node is not None and node.tag == qn("w:p"):
                text = "".join(t.text or "" for t in node.findall(".//" + qn("w:t"))).strip()
                if text:
                    near_text.append(text)
        if any(re.match(pattern, text, re.I) for text in near_text for pattern in table_patterns):
            continue
        anchor_index = 0
        for idx, para in enumerate(doc.paragraphs):
            if prev is not None and para._p is prev:
                anchor_index = idx
                break
        table_index = _next_unused(_find_nearby_number(doc.paragraphs, anchor_index, r"表\s*(\d+)"), used_table_numbers)
        cn_para = _new_paragraph_before_table(table, TEST_CAPTIONS["table_cn"].format(index=table_index))
        en_para = _new_paragraph_before_table(table, TEST_CAPTIONS["table_en"].format(index=table_index))
        cn_para.paragraph_format.keep_with_next = True
        en_para.paragraph_format.keep_with_next = True
        if blocks.get("caption_table"):
            _apply_paragraph_style(cn_para, blocks["caption_table"])
            _apply_paragraph_style(en_para, blocks["caption_table"])



def _ensure_optional_blocks(doc: Document, blocks: Dict[str, Dict[str, Any]], template_blocks: Dict[str, Any]) -> None:
    references_title = _find_paragraph_starting(doc, TEXT["references"])
    references_index = _paragraph_index(doc, references_title) if references_title is not None else None
    footer_para = _find_paragraph_starting(doc, "\u6536\u7a3f\u65e5\u671f")
    author_bio_title = _find_paragraph_starting(doc, TEXT["author_bio"])

    fund_exists = any(_text(p.text).startswith("\u57fa\u91d1\u9879\u76ee") for p in doc.paragraphs)
    if not fund_exists and blocks.get("fund"):
        anchor = doc.paragraphs[references_index - 1] if references_index is not None and references_index > 0 else doc.paragraphs[-1]
        _insert_block_after(anchor, TEST_OPTIONAL_BLOCKS["fund"], blocks.get("fund"), template_blocks.get("fund"))

    if references_title is None:
        return

    refs_started = False
    last_ref_para = references_title
    for paragraph in doc.paragraphs:
        text = _text(paragraph.text)
        if paragraph._p is references_title._p:
            refs_started = True
            continue
        if not refs_started:
            continue
        if text.startswith(TEXT["author_bio"]) or text.startswith("\u6536\u7a3f\u65e5\u671f"):
            break
        if _is_reference(text):
            last_ref_para = paragraph

    if footer_para is None:
        footer_para = _insert_block_after(last_ref_para, TEST_OPTIONAL_BLOCKS["footer"], blocks.get("footer"), template_blocks.get("footer"))
    if author_bio_title is None:
        author_bio_title = _insert_block_after(footer_para, TEST_OPTIONAL_BLOCKS["author_bio_title"], blocks.get("author_bio_title"), template_blocks.get("author_bio_title"))

    bio_entries = []
    bio_started = False
    for paragraph in doc.paragraphs:
        text = _text(paragraph.text)
        if author_bio_title is not None and paragraph._p is author_bio_title._p:
            bio_started = True
            continue
        if not bio_started:
            continue
        if re.match(r"^[^()（）]{2,20}[（(]", text):
            bio_entries.append(paragraph)
    if not bio_entries and blocks.get("author_bio_entry"):
        first = _insert_block_after(author_bio_title, TEST_OPTIONAL_BLOCKS["author_bio_entry_1"], blocks.get("author_bio_entry"), template_blocks.get("author_bio_entry"))
        second = _insert_block_after(first, TEST_OPTIONAL_BLOCKS["author_bio_entry_2"], blocks.get("author_bio_entry"), template_blocks.get("author_bio_entry"))
        _insert_block_after(second, TEST_OPTIONAL_BLOCKS["author_bio_entry_3"], blocks.get("author_bio_entry"), template_blocks.get("author_bio_entry"))


def _index_map(doc: Document) -> Dict[str, int | None]:
    abstract_cn = _find_paragraph_starting(doc, TEXT["abstract"])
    keywords_cn = _find_paragraph_starting(doc, (TEXT["keywords"], TEXT["keywords_alt"]))
    intro = _find_paragraph_starting(doc, "引言")
    title_en = _find_paragraph_starting(doc, "health protection system innovation")
    abstract_en = _find_paragraph_starting(doc, "abstract")
    return {
        "abstract_cn": _paragraph_index(doc, abstract_cn) if abstract_cn is not None else None,
        "keywords_cn": _paragraph_index(doc, keywords_cn) if keywords_cn is not None else None,
        "intro": _paragraph_index(doc, intro) if intro is not None else None,
        "title_en": _paragraph_index(doc, title_en) if title_en is not None else None,
        "abstract_en": _paragraph_index(doc, abstract_en) if abstract_en is not None else None,
    }


def _set_twips_attr(element, attr: str, pt_value) -> None:
    if pt_value is None:
        return
    element.set(qn(attr), str(int(round(float(pt_value) * 20))))


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _apply_section_layout_to_sectpr(sect_pr, layout: Dict[str, Any], continuous: bool | None = None) -> None:
    if sect_pr is None or not layout:
        return

    sect_type = sect_pr.find(qn("w:type"))
    if continuous is True:
        if sect_type is None:
            sect_type = OxmlElement("w:type")
            sect_pr.insert(0, sect_type)
        sect_type.set(qn("w:val"), "continuous")
    elif continuous is False and sect_type is not None:
        sect_pr.remove(sect_type)

    pg_sz = _ensure_child(sect_pr, "w:pgSz")
    _set_twips_attr(pg_sz, "w:w", layout.get("page_width"))
    _set_twips_attr(pg_sz, "w:h", layout.get("page_height"))

    pg_mar = _ensure_child(sect_pr, "w:pgMar")
    _set_twips_attr(pg_mar, "w:top", layout.get("top_margin"))
    _set_twips_attr(pg_mar, "w:bottom", layout.get("bottom_margin"))
    _set_twips_attr(pg_mar, "w:left", layout.get("left_margin"))
    _set_twips_attr(pg_mar, "w:right", layout.get("right_margin"))
    _set_twips_attr(pg_mar, "w:header", layout.get("header_distance"))
    _set_twips_attr(pg_mar, "w:footer", layout.get("footer_distance"))

    cols = _ensure_child(sect_pr, "w:cols")
    columns_num = layout.get("columns_num")
    columns_space = layout.get("columns_space_pt")
    if columns_num is not None:
        cols.set(qn("w:num"), str(int(columns_num)))
    if columns_space is not None:
        cols.set(qn("w:space"), str(int(round(float(columns_space) * 20))))


def _find_section_break_index(doc: Document) -> int | None:
    for index, paragraph in enumerate(doc.paragraphs):
        text = _text(paragraph.text)
        if not text:
            continue
        if text == "引言":
            return index
        if _is_heading_1(text):
            return index
    return None


def _apply_section_breaks(doc: Document, rules: Dict[str, Any]) -> None:
    section_layouts = rules.get("section_layouts", [])
    if len(section_layouts) < 2 or not doc.sections:
        return

    first_layout = section_layouts[0]
    remaining_layout = section_layouts[1]
    _apply_section_layout_to_sectpr(doc.sections[0]._sectPr, first_layout, continuous=False)

    break_index = _find_section_break_index(doc)
    if break_index is None or break_index == 0:
        return

    anchor = doc.paragraphs[break_index - 1]
    anchor_ppr = anchor._p.get_or_add_pPr()
    existing_anchor = anchor_ppr.find(qn("w:sectPr"))
    if existing_anchor is not None:
        anchor_ppr.remove(existing_anchor)

    body_sect_pr = doc.sections[-1]._sectPr
    anchor_sect_pr = deepcopy(body_sect_pr)
    _apply_section_layout_to_sectpr(anchor_sect_pr, first_layout, continuous=False)
    anchor_ppr.append(anchor_sect_pr)

    _apply_section_layout_to_sectpr(body_sect_pr, remaining_layout, continuous=True)


def _apply_template_headers_footers(doc: Document, template_doc: Document | None) -> None:
    if template_doc is None or not doc.sections or not template_doc.sections:
        return
    template_sections = list(template_doc.sections)
    for index, section in enumerate(doc.sections):
        template_section = template_sections[min(index, len(template_sections) - 1)]
        section.different_first_page_header_footer = template_section.different_first_page_header_footer
        # Do not mirror linked-to-previous flags directly. Rebuilding those
        # relations across packages can leave python-docx with broken header
        # parts after save/reload. Instead, materialize local definitions and
        # copy the template content into each section.
        _rewrite_story_from_template(template_section.header, section.header)
        _rewrite_story_from_template(template_section.first_page_header, section.first_page_header)
        _rewrite_story_from_template(template_section.footer, section.footer)
        _rewrite_story_from_template(template_section.first_page_footer, section.first_page_footer)


def _apply_story_rule(story, rule: Dict[str, Any] | None) -> None:
    if not rule:
        return
    for paragraph in story.paragraphs:
        if _text(paragraph.text):
            _apply_paragraph_style(paragraph, rule)


def _apply_header_footer_rules(doc: Document, blocks: Dict[str, Any]) -> None:
    for section in doc.sections:
        _apply_story_rule(section.header, blocks.get("header"))
        _apply_story_rule(section.first_page_header, blocks.get("header"))
        _apply_story_rule(section.footer, blocks.get("footer"))
        _apply_story_rule(section.first_page_footer, blocks.get("footer"))


def format_doc(input_docx: Path, rules_json: Path, output_docx: Path) -> None:
    rules = json.loads(rules_json.read_text(encoding="utf-8"))
    doc = Document(str(input_docx))
    template_doc = None
    template_blocks: Dict[str, Any] = {}
    template_path = rules.get("source", {}).get("template_path")
    if template_path:
        template_candidate = Path(template_path)
        if template_candidate.exists():
            template_doc = Document(str(template_candidate))
            detected = detect_blocks(template_doc)
            for role, paragraphs in detected.items():
                filtered = filter_instructional_paragraphs(paragraphs, role)
                if filtered:
                    template_blocks[role] = filtered[0]
                elif paragraphs:
                    template_blocks[role] = paragraphs[0]

    if doc.sections:
        section = doc.sections[0]
        layout = rules.get("page_layout", {})
        if layout.get("page_width") is not None:
            section.page_width = Pt(layout["page_width"])
        if layout.get("page_height") is not None:
            section.page_height = Pt(layout["page_height"])
        if layout.get("top_margin") is not None:
            section.top_margin = Pt(layout["top_margin"])
        if layout.get("bottom_margin") is not None:
            section.bottom_margin = Pt(layout["bottom_margin"])
        if layout.get("left_margin") is not None:
            section.left_margin = Pt(layout["left_margin"])
        if layout.get("right_margin") is not None:
            section.right_margin = Pt(layout["right_margin"])
        if layout.get("header_distance") is not None:
            section.header_distance = Pt(layout["header_distance"])
        if layout.get("footer_distance") is not None:
            section.footer_distance = Pt(layout["footer_distance"])

    blocks = rules.get("blocks", {})
    _sync_template_styles(doc, template_doc, template_blocks)
    _ensure_front_matter(doc, blocks)
    _apply_section_breaks(doc, rules)
    _apply_template_headers_footers(doc, template_doc)
    _apply_header_footer_rules(doc, blocks)
    _ensure_caption_placeholders(doc, blocks)
    _ensure_optional_blocks(doc, blocks, template_blocks)
    idx = _index_map(doc)

    title_done = False
    refs_started = False
    author_bio_started = False
    for para_index, paragraph in enumerate(doc.paragraphs):
        text = _text(paragraph.text)
        lower_text = text.lower()
        if not text:
            continue
        if not title_done and blocks.get("title") and len(text) > 8 and not lower_text.startswith("doi:"):
            _apply_template_block(paragraph, blocks["title"], template_blocks.get("title"))
            title_done = True
            continue
        if (
            blocks.get("authors")
            and idx["abstract_cn"] is not None
            and title_done
            and para_index < idx["abstract_cn"]
            and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
            and not text.startswith("(")
            and len(text) < 60
        ):
            _apply_template_block(paragraph, blocks["authors"], template_blocks.get("authors"))
            continue
        if (
            blocks.get("affiliations")
            and idx["abstract_cn"] is not None
            and title_done
            and para_index < idx["abstract_cn"]
            and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
            and text.startswith("(")
        ):
            _apply_template_block(paragraph, blocks["affiliations"], template_blocks.get("affiliations"))
            continue
        if text.startswith(TEXT["abstract"]) and blocks.get("abstract_cn"):
            _apply_template_block(paragraph, blocks["abstract_cn"], template_blocks.get("abstract_cn"))
            continue
        if text.startswith((TEXT["keywords"], TEXT["keywords_alt"])) and blocks.get("keywords_cn"):
            _apply_template_block(paragraph, blocks["keywords_cn"], template_blocks.get("keywords_cn"))
            continue
        if (
            idx["keywords_cn"] is not None
            and idx["intro"] is not None
            and idx["keywords_cn"] < para_index < idx["intro"]
            and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
            and blocks.get("title_en")
            and para_index == idx["title_en"]
        ):
            _apply_template_block(paragraph, blocks["title_en"], template_blocks.get("title_en"))
            continue
        if (
            idx["keywords_cn"] is not None
            and idx["intro"] is not None
            and idx["keywords_cn"] < para_index < idx["intro"]
            and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
            and blocks.get("authors_en")
            and idx["title_en"] is not None
            and para_index == idx["title_en"] + 1
            and not text.startswith("(")
        ):
            _apply_template_block(paragraph, blocks["authors_en"], template_blocks.get("authors_en"))
            continue
        if (
            idx["keywords_cn"] is not None
            and idx["intro"] is not None
            and idx["keywords_cn"] < para_index < idx["intro"]
            and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
            and blocks.get("affiliations_en")
            and idx["title_en"] is not None
            and para_index == idx["title_en"] + 2
            and text.startswith("(")
        ):
            _apply_template_block(paragraph, blocks["affiliations_en"], template_blocks.get("affiliations_en"))
            continue
        if lower_text.startswith("abstract") and blocks.get("abstract_en"):
            _apply_template_block(paragraph, blocks["abstract_en"], template_blocks.get("abstract_en"))
            continue
        if lower_text.startswith("key words") and blocks.get("keywords_en"):
            _apply_template_block(paragraph, blocks["keywords_en"], template_blocks.get("keywords_en"))
            continue
        if text.startswith("\u57fa\u91d1\u9879\u76ee") and blocks.get("fund"):
            _apply_template_block(paragraph, blocks["fund"], template_blocks.get("fund"))
            continue
        if text.startswith(TEXT["references"]):
            refs_started = True
            if blocks.get("references_title"):
                _apply_paragraph_style(paragraph, blocks["references_title"])
            continue
        if text.startswith("\u6536\u7a3f\u65e5\u671f") and blocks.get("footer"):
            _apply_template_block(paragraph, blocks["footer"], template_blocks.get("footer"))
            continue
        if text.startswith(TEXT["author_bio"]):
            author_bio_started = True
            if blocks.get("author_bio_title"):
                _apply_template_block(paragraph, blocks["author_bio_title"], template_blocks.get("author_bio_title"))
            continue
        if refs_started and _is_reference(text) and blocks.get("reference_entry"):
            _apply_paragraph_style(paragraph, blocks["reference_entry"])
            continue
        if author_bio_started and blocks.get("author_bio_entry") and re.match(r"^[^()\uFF08\uFF09]{2,20}[\uFF08(]", text):
            _apply_template_block(paragraph, blocks["author_bio_entry"], template_blocks.get("author_bio_entry"))
            continue
        if (TEXT["figure_caption"] in text or re.match(r"^(\u56fe|Figure|Fig\.)", text, re.I)) and blocks.get("caption_figure"):
            _apply_template_block(paragraph, blocks["caption_figure"], template_blocks.get("caption_figure"))
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            continue
        if (TEXT["table_caption"] in text or re.match(r"^(\u8868|Table)", text, re.I)) and blocks.get("caption_table"):
            _apply_template_block(paragraph, blocks["caption_table"], template_blocks.get("caption_table"))
            paragraph.paragraph_format.keep_with_next = True
            continue
        if is_heading_1_paragraph(paragraph) and blocks.get("heading_1"):
            _apply_template_block(paragraph, blocks["heading_1"], template_blocks.get("heading_1"))
            continue
        if is_heading_2_paragraph(paragraph) and blocks.get("heading_2"):
            _apply_template_block(paragraph, blocks["heading_2"], template_blocks.get("heading_2"))
            continue
        if blocks.get("body") and len(text) >= 18 and paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            _apply_template_block(paragraph, blocks["body"], template_blocks.get("body"))

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))


def main() -> int:
    parser = argparse.ArgumentParser(description="Apply conservative formatting rules to a manuscript .docx")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--rules", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    format_doc(args.input, args.rules, args.output)
    print(f"Formatted document written to {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
