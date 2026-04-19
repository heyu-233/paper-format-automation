from __future__ import annotations

import sys
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT / "skills" / "paper-format-automation" / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import docx_rule_utils as utils


class DocxRuleUtilsTests(unittest.TestCase):
    def test_instructional_text_detection_matches_formatting_guidance(self) -> None:
        text = "正文使用宋体小四号，段后 6 磅，行距 1.5 倍，图题使用黑体。"
        self.assertTrue(utils.is_instructional_text(text))

    def test_heading_detection_distinguishes_body_and_subsections(self) -> None:
        self.assertTrue(utils.is_heading_1_text("1 研究背景"))
        self.assertFalse(utils.is_heading_1_text("1.1 研究背景"))
        self.assertTrue(utils.is_heading_2_text("2.3 方法设计"))
        self.assertFalse(utils.is_heading_2_text("研究背景"))

    def test_detect_blocks_finds_front_matter_roles(self) -> None:
        doc = Document()
        doc.add_paragraph("基于模板的期刊格式自动化")
        author_para = doc.add_paragraph("张三1，李四2")
        author_para.alignment = 1
        affiliation_para = doc.add_paragraph("(1. 示例单位，北京 100000；2. 示例单位，上海 200000)")
        affiliation_para.alignment = 1
        doc.add_paragraph("摘要：这里是中文摘要内容。")
        doc.add_paragraph("关键词：格式自动化；模板对齐")
        title_en = doc.add_paragraph("Template-Driven Journal Formatting Automation")
        title_en.alignment = 1
        authors_en = doc.add_paragraph("ZHANG San, LI Si")
        authors_en.alignment = 1
        affiliation_en = doc.add_paragraph("(1. Example Institute, Beijing 100000, China)")
        affiliation_en.alignment = 1
        doc.add_paragraph("Abstract: English abstract.")
        doc.add_paragraph("Key words: formatting; automation")
        doc.add_paragraph("1 研究背景")
        doc.add_paragraph("这是足够长的正文段落，用来确保正文识别逻辑可以抓到实际内容。")

        blocks = utils.detect_blocks(doc)

        self.assertEqual(blocks["title"][0].text, "基于模板的期刊格式自动化")
        self.assertEqual(blocks["authors"][0].text, "张三1，李四2")
        self.assertEqual(blocks["affiliations"][0].text, "(1. 示例单位，北京 100000；2. 示例单位，上海 200000)")
        self.assertEqual(blocks["title_en"][0].text, "Template-Driven Journal Formatting Automation")
        self.assertEqual(blocks["heading_1"][0].text, "1 研究背景")
        self.assertTrue(blocks["body"])


if __name__ == "__main__":
    unittest.main()
