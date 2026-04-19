from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT / "skills" / "paper-format-automation" / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import check_manuscript
import docx_rule_utils as utils
import format_manuscript
from extract_template_rules import build_rules


def _set_run_fonts(run, *, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    east_asia: str = "宋体",
    ascii_font: str = "Times New Roman",
    size_pt: float = 10.5,
    bold: bool = False,
    first_line_indent_pt: float | None = None,
    line_spacing: float | None = None,
) -> None:
    paragraph = doc.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    if first_line_indent_pt is not None:
        paragraph.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing
    run = paragraph.add_run(text)
    _set_run_fonts(run, east_asia=east_asia, ascii_font=ascii_font, size_pt=size_pt, bold=bold)


def _make_template_doc(path: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(90)
    section.bottom_margin = Pt(90)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    _add_paragraph(doc, "基于模板的期刊格式自动化", alignment=WD_ALIGN_PARAGRAPH.CENTER, east_asia="黑体", size_pt=16, bold=True)
    _add_paragraph(doc, "张三1，李四2", alignment=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12)
    _add_paragraph(doc, "(1. 示例单位，北京 100000；2. 示例单位，上海 200000)", alignment=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5)
    _add_paragraph(doc, "摘要：这是模板摘要内容。", size_pt=10.5)
    _add_paragraph(doc, "关键词：格式自动化；模板对齐", size_pt=10.5)
    _add_paragraph(doc, "Template-Driven Journal Formatting Automation", alignment=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12, bold=True)
    _add_paragraph(doc, "ZHANG San, LI Si", alignment=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5)
    _add_paragraph(doc, "(1. Example Institute, Beijing 100000, China)", alignment=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5)
    _add_paragraph(doc, "Abstract: This is the template abstract.", size_pt=10.5)
    _add_paragraph(doc, "Key words: formatting; automation", size_pt=10.5)
    _add_paragraph(doc, "1 研究背景", east_asia="黑体", size_pt=12, bold=True)
    _add_paragraph(doc, "这是模板中的正文段落，用来沉淀正文样式规则并驱动格式化测试。", first_line_indent_pt=21, line_spacing=1.5)

    doc.save(str(path))
    return path


def _make_manuscript_doc(path: Path, *, include_english_front_matter: bool) -> Path:
    doc = Document()
    _add_paragraph(doc, "基于模板的期刊格式自动化", alignment=WD_ALIGN_PARAGRAPH.LEFT, east_asia="宋体", size_pt=12, bold=False)
    _add_paragraph(doc, "张三1，李四2", alignment=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5)
    _add_paragraph(doc, "(1. 示例单位，北京 100000；2. 示例单位，上海 200000)", alignment=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5)
    _add_paragraph(doc, "摘要：这是待检查稿件的摘要。", size_pt=10.5)
    _add_paragraph(doc, "关键词：自动化；测试", size_pt=10.5)
    if include_english_front_matter:
        _add_paragraph(doc, "Template-Driven Journal Formatting Automation", alignment=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5)
        _add_paragraph(doc, "ZHANG San, LI Si", alignment=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5)
        _add_paragraph(doc, "(1. Example Institute, Beijing 100000, China)", alignment=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5)
        _add_paragraph(doc, "Abstract: This is the manuscript abstract.", size_pt=10.5)
        _add_paragraph(doc, "Key words: testing; automation", size_pt=10.5)
    _add_paragraph(doc, "1 研究背景", east_asia="宋体", size_pt=10.5, bold=False)
    _add_paragraph(doc, "这是待格式化稿件中的正文段落，它故意保留了不符合模板的缩进和字号。", first_line_indent_pt=0, line_spacing=1.0)

    doc.save(str(path))
    return path


class EndToEndDocxFlowTests(unittest.TestCase):
    def test_check_manuscript_cli_writes_actionable_diff_outputs(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            template_path = _make_template_doc(tmp_path / "template.docx")
            manuscript_path = _make_manuscript_doc(tmp_path / "manuscript.docx", include_english_front_matter=True)
            rules_path = tmp_path / "template_rules.json"
            report_path = tmp_path / "diff_report.json"
            markdown_path = tmp_path / "review-report.md"

            rules = build_rules(template_path)
            rules_path.write_text(json.dumps(rules, ensure_ascii=False, indent=2), encoding="utf-8")

            argv = [
                "check_manuscript.py",
                str(manuscript_path),
                str(rules_path),
                "-o",
                str(report_path),
                "--markdown",
                str(markdown_path),
            ]
            with patch.object(sys, "argv", argv):
                rc = check_manuscript.main()

            self.assertEqual(rc, 0)
            report = json.loads(report_path.read_text(encoding="utf-8"))
            title_check = next(item for item in report["block_checks"] if item["role"] == "title")
            body_check = next(item for item in report["block_checks"] if item["role"] == "body")

            self.assertEqual(report["page_layout"]["status"], "auto_fix_candidate")
            self.assertEqual(title_check["status"], "auto_fix_candidate")
            self.assertEqual(body_check["status"], "auto_fix_candidate")
            self.assertTrue(any(item["field"] == "top_margin" for item in report["page_layout"]["mismatches"]))

            markdown = markdown_path.read_text(encoding="utf-8")
            self.assertIn("## Block Checks", markdown)
            self.assertIn("正文 (body) - auto_fix_candidate", markdown)

    def test_format_manuscript_cli_applies_rules_and_reduces_mismatches(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            template_path = _make_template_doc(tmp_path / "template.docx")
            manuscript_path = _make_manuscript_doc(tmp_path / "manuscript.docx", include_english_front_matter=False)
            rules_path = tmp_path / "template_rules.json"
            formatted_path = tmp_path / "formatted.docx"

            rules = build_rules(template_path)
            rules_path.write_text(json.dumps(rules, ensure_ascii=False, indent=2), encoding="utf-8")

            before_report = check_manuscript.build_report(manuscript_path, rules_path)
            before_auto_fix = before_report["summary"]["auto_fix_candidate"]

            argv = [
                "format_manuscript.py",
                "--input",
                str(manuscript_path),
                "--rules",
                str(rules_path),
                "--output",
                str(formatted_path),
            ]
            with patch.object(sys, "argv", argv):
                rc = format_manuscript.main()

            self.assertEqual(rc, 0)
            formatted_doc = Document(str(formatted_path))
            texts = [paragraph.text for paragraph in formatted_doc.paragraphs if paragraph.text.strip()]
            title_snapshot = utils.paragraph_snapshot(formatted_doc.paragraphs[0])
            body_paragraph = next(paragraph for paragraph in formatted_doc.paragraphs if "待格式化稿件中的正文段落" in paragraph.text)
            body_snapshot = utils.paragraph_snapshot(body_paragraph)
            after_report = check_manuscript.build_report(formatted_path, rules_path)

            self.assertIn(format_manuscript.TEST_FRONT_MATTER["title_en"], texts)
            self.assertIn(format_manuscript.TEST_FRONT_MATTER["abstract_en"], texts)
            self.assertEqual(title_snapshot["alignment"], "center")
            self.assertEqual(title_snapshot["font_size_pt"], 16.0)
            self.assertTrue(title_snapshot["bold"])
            self.assertEqual(body_snapshot["first_line_indent_pt"], 21.0)
            self.assertEqual(body_snapshot["line_spacing_mode"], "multiple")
            self.assertEqual(body_snapshot["line_spacing_value"], 1.5)
            self.assertLess(after_report["summary"]["auto_fix_candidate"], before_auto_fix)


if __name__ == "__main__":
    unittest.main()
